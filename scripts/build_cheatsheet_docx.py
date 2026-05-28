#!/usr/bin/env python3
"""OpenClaw RL cheatsheet — Word version.

Sections:
  1. Cover + master pre-submission checklist
  2. Prompt cheatsheet + common errors
  3. Prompt-authoring FLOWCHART (8-step decision tree)
  4. Rubric cheatsheet + common errors
  5. Standalone Rubric CHECKLIST with scoring
  6. Silver trajectory cheatsheet + common errors
  7. Tests cheatsheet + common errors
  8. 21 QC Dimensions + cardinal rules
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

OUT = Path("/Users/vishal.kushwaha/Documents/Openclaw_RL copy/OpenClaw_Cheatsheet.docx")

NAVY = RGBColor(0x1F, 0x38, 0x64)
MID = RGBColor(0x2E, 0x53, 0x95)
GREY = RGBColor(0x44, 0x44, 0x44)
RED = RGBColor(0x9C, 0x22, 0x22)
GREEN = RGBColor(0x2E, 0x7D, 0x32)


def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = MID


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = GREY


def parse_inline(text):
    out, i, n = [], 0, len(text)
    while i < n:
        if text.startswith("**", i):
            j = text.find("**", i + 2)
            if j == -1:
                out.append((text[i:], False, False)); break
            out.append((text[i + 2:j], True, False)); i = j + 2
        elif text[i] == "`":
            j = text.find("`", i + 1)
            if j == -1:
                out.append((text[i:], False, False)); break
            out.append((text[i + 1:j], False, True)); i = j + 1
        else:
            j = i
            while j < n and text[j] != "*" and text[j] != "`":
                j += 1
            out.append((text[i:j], False, False)); i = j
    return out


def add_body(doc, text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    for txt, bold, italic in parse_inline(text):
        r = p.add_run(txt); r.bold = bold; r.italic = italic; r.font.size = Pt(size)


def add_bullet(doc, text, check=False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.2)
    if check:
        r0 = p.add_run("☐ "); r0.font.size = Pt(10)
    for txt, bold, italic in parse_inline(text):
        r = p.add_run(txt); r.bold = bold; r.italic = italic; r.font.size = Pt(10)


def add_table(doc, rows, col_widths=None, alt_rows=True, hdr_bg="1F3864",
              cell_colors=None):
    """cell_colors: dict {(row, col): hex} for per-cell shading."""
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid"
    if col_widths:
        for col, w in enumerate(col_widths):
            for cell in table.columns[col].cells:
                cell.width = Inches(w)
    cell_colors = cell_colors or {}
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for para in cell.paragraphs:
                para.clear()
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            r.font.size = Pt(8) if i > 0 else Pt(9)
            if i == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_bg(cell, hdr_bg)
            elif (i, j) in cell_colors:
                set_cell_bg(cell, cell_colors[(i, j)])
            elif alt_rows and i % 2 == 0:
                set_cell_bg(cell, "F2F2F2")
    doc.add_paragraph()


def page_break(doc):
    doc.add_page_break()


# ============ FLOWCHART ============
def add_flowchart_step(doc, step_no, question, yes_to, no_action, *, is_start=False, is_end=False):
    """Render a single flowchart node as a styled mini-table."""
    if is_start:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Light Grid"
        cell = tbl.cell(0, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.MIDDLE
        for para in cell.paragraphs: para.clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("▶  START — Prompt draft complete?")
        r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, "2E7D32")
        _arrow(doc)
        return

    if is_end:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Light Grid"
        cell = tbl.cell(0, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.MIDDLE
        for para in cell.paragraphs: para.clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("■  READY — Prompt passes all 8 gates. Submit.")
        r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, "1F3864")
        return

    # 3-column row: step | question/yes | no/action
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Light Grid"
    # step
    cell = tbl.cell(0, 0); cell.width = Inches(0.4)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.MIDDLE
    for para in cell.paragraphs: para.clear()
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Q{step_no}"); r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_bg(cell, "1F3864")
    # question
    cell = tbl.cell(0, 1); cell.width = Inches(4.6)
    for para in cell.paragraphs: para.clear()
    p = cell.paragraphs[0]
    r = p.add_run(question); r.bold = True; r.font.size = Pt(9.5)
    p2 = cell.add_paragraph()
    r = p2.add_run(f"✓ YES → {yes_to}")
    r.font.size = Pt(8.5); r.font.color.rgb = GREEN
    # no action
    cell = tbl.cell(0, 2); cell.width = Inches(2.6)
    set_cell_bg(cell, "FFE5E5")
    for para in cell.paragraphs: para.clear()
    p = cell.paragraphs[0]
    r = p.add_run("✗ NO → STOP, FIX:"); r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = RED
    p2 = cell.add_paragraph()
    r = p2.add_run(no_action); r.font.size = Pt(8.5)
    _arrow(doc)


def _arrow(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("▼")
    r.font.size = Pt(11); r.bold = True; r.font.color.rgb = NAVY


# ============ MAIN BUILD ============
def build():
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ========== COVER ==========
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("OpenClaw RL — QC Cheatsheet")
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("For attempters & reviewers  |  Prompt + Rubric + Tests + Flowchart + Checklist")
    r.italic = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x5E, 0x5E, 0x5E)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Distilled from 382 universe-grounded audits across 5 batches (test_v3, L12, test4, L11, L4)")
    r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0x5E, 0x5E, 0x5E)

    add_h2(doc, "What's in this cheatsheet")
    add_table(doc, [
        ["Section", "Purpose"],
        ["1. Prompt cheatsheet + common errors",       "Pitfalls + quick hacks for the opening prompt"],
        ["2. Prompt-authoring FLOWCHART", "8-step decision tree to walk through before submission"],
        ["3. Rubric cheatsheet + common errors",       "Top defects + quick hacks for rubric authors"],
        ["4. Rubric CHECKLIST (scored)", "Line-item gate with thresholds — score your rubric"],
        ["5. Silver trajectory cheatsheet + common errors", "What kills a Silver — defects we caught in audits"],
        ["6. Tests cheatsheet + common errors",        "Verifier defects + how to avoid them"],
        ["7. 21 QC Dimensions + cardinal rules", "What auditors grade and how"],
    ], col_widths=[2.2, 5.0])

    add_h2(doc, "Headline failure rates")
    add_table(doc, [
        ["Batch", "Tasks", "% Score 2", "% Score 4-5", "Top failure mode"],
        ["test_v3",   "4",   "75%",  "25%",  "Aggressive G5 / DO-leak flagging"],
        ["L12",       "87",  "56%",  "37%",  "Rubric overfits DO-only section headers"],
        ["test4",     "20",  "40%",  "55%",  "Rubric weight-sign inversions"],
        ["L11",       "103", "55%",  "36%",  "Rubric / prompt contradiction"],
        ["L4 (att.)", "168", "47%",  "44%",  "CB justifications defend unrequested rubric"],
    ], col_widths=[1.0, 0.7, 0.85, 0.85, 3.0])

    add_h2(doc, "Master pre-submission checklist (1 minute)")
    add_table(doc, [
        ["#", "PROMPT", "RUBRIC", "TESTS"],
        ["1",
         "Every artifact name and section header you grade is in the opening_prompt OR a workspace file you tell the model to read",
         "Each criterion is defensible against opening_prompt alone — not against desired_outcome",
         "No hardcoded number contradicts the universe snapshot"],
        ["2",
         "If you say \"do not produce X yet\" or \"keep in chat,\" the rubric does NOT grade producing X",
         "Polarity matches intent: positive weight rewards good; negative weight penalizes bad",
         "No EXPECTED_COUNTS hardcoded that contradict snapshots.json"],
        ["3",
         "Required files referenced are real (rule.md, policy.md, output_spec.md, attached PDFs/images)",
         "Atomicity: one constraint per criterion. Do not bundle \"X AND Y AND Z\"",
         "No test requires an artifact the prompt explicitly forbids"],
        ["4",
         "Safety boundary, tool channels, and recipients are explicit",
         "Safety tasks carry at least one negative-weight criterion",
         "Date comparisons include time component, not just first 10 chars"],
        ["5",
         "Numeric thresholds and dates are defined or derivable from universe data",
         "Justifications match criterion titles, no \"correct it is\" boilerplate",
         "Filename-locked tests are only used when the filename is in opening_prompt"],
    ], col_widths=[0.25, 2.1, 2.1, 2.1])

    page_break(doc)

    # ========== 1. PROMPT ==========
    add_h1(doc, "1. PROMPT cheatsheet")
    add_body(doc,
             "The opening prompt is the contract. Everything the rubric grades and every test asserts must trace back to "
             "the opening prompt OR a workspace file the model is explicitly told to read. If your rubric needs a section "
             "header the prompt doesn't mention, the prompt is incomplete (not the model wrong).")

    add_h2(doc, "Pre-submission checklist")
    for item in [
        "Every deliverable filename appears in the opening prompt or in a referenced workspace file",
        "Every required section heading or label vocabulary is named in the opening prompt",
        "Required tool channels (email, slack, calendar, fintrack, airtable, etc.) are named explicitly",
        "Forbidden actions are stated explicitly (\"do not send\", \"do not change records\", \"do not write files\")",
        "Multi-turn gates (\"first review, then produce\") are clear about which turn is the final deliverable",
        "Numeric thresholds (counts, percentages, dates, windows) are stated or derivable from universe data",
        "Privacy/safety boundaries (who can receive what) are explicit",
        "Source-of-truth files (rule.md, policy.md, policy.png, output_spec.md, *.pdf) are named for the model to read",
        "PII the model may legitimately see vs surface is distinguished",
        "Universe data the rubric depends on is reachable (no fabricated services like \"FinTrack\" when only fintrack exists)",
    ]:
        add_bullet(doc, item, check=True)

    add_h2(doc, "Top 14 prompt errors (every one caught in real audits)")
    for title, body in [
        ("Prompt forbids files, rubric demands files",
         "Prompt: \"keep in chat\". Rubric: grades creation of report.docx. ~22 L4 tasks. Remove one constraint."),
        ("Prompt says \"do not produce X yet\"",
         "Multi-turn plan-first prompt; rubric still penalizes missing final-deck-now (case `6a04c1b0...57a1`). Pick one mode."),
        ("Rubric demands artifacts named only in desired_outcome",
         "Prompt: \"a memory file and a plan\". Rubric: exact filenames MEMORY.md, financial_health_report.md, financial_action_plan.md."),
        ("Universe service name is fabricated",
         "Rubric references \"FinTrack\" but universe has fintrack. Verify service names against the actual universe."),
        ("Prompt contradicts test/rubric",
         "Prompt: \"don't change records.\" Tests: \"MEMORY.md must contain narrative X.\" The model literally cannot pass."),
        ("Filename in policy.png/PDF not surfaced",
         "Required filename colt_reply_draft.md lives only in policy.png. State it in the prompt body too."),
        ("Workspace file referenced but missing",
         "Prompt directs read of memory_seed.md; universe lacks it. Model halts; rubric punishes the halt."),
        ("Casual phrasing graded as literal string",
         "Prompt: \"a clean one-page brief.\" Rubric: exact 36-word/6-line cap. Codify in the prompt or relax the rubric."),
        ("Referenced tool/service doesn't exist in the universe",
         "Prompt says \"check FinTrack\" (capital F); universe only has `fintrack`. Saw on task `5835`. "
         "Verify all named tools/services against the actual universe before submitting."),
        ("Conflicting instructions in the same prompt",
         "\"Be thorough and include every detail\" + \"Keep it under 50 words.\" The model satisfies one and fails "
         "the other; the rubric punishes whichever it cares about. Pick one and only one."),
        ("Missing units or time zones",
         "\"Schedule for 9 AM\" with no time zone; \"$100\" with no currency; \"a 5-minute walk\" with no distance. "
         "Saw the 9 AM / 14:00Z mismatch on andrew_mitchell tasks — model wrote `T090000` because the prompt didn't disambiguate."),
        ("Implicit \"standard format\" with no template",
         "\"Generate the standard quarterly report\" without an attached template and no rubric criterion defining "
         "\"standard.\" The model invents a format; the rubric punishes the invention."),
        ("Source-of-truth file referenced but model isn't told to read it",
         "Prompt says \"see attached policy.md\"; doesn't say \"read it first.\" Model may skip. Use explicit imperatives: "
         "\"Read `policy.md` BEFORE producing any output.\""),
    ]:
        add_h3(doc, title); add_body(doc, body)

    add_h2(doc, "Quick hacks for attempters")
    for h in [
        "Before writing the rubric, list every concrete word/number you plan to grade. Each must appear in the prompt body or a referenced workspace file.",
        "Use **must** / **must not** / **do not** for hard constraints — reviewers and models parse these consistently.",
        "Multi-turn? Label the final-deliverable turn explicitly: \"In your final response, produce X.\"",
        "Test that any referenced workspace file (rule.md, policy.png) is actually placed in the universe.",
        "If universe data conflicts with rubric assumptions, fix the prompt or the universe — never patch with a rubric carve-out.",
    ]:
        add_bullet(doc, h)

    page_break(doc)

    # ========== 2. PROMPT FLOWCHART ==========
    add_h1(doc, "2. Prompt-authoring FLOWCHART")
    add_body(doc,
             "Walk through these 8 gates before submitting. Stop and fix any time you answer NO. "
             "Every red gate maps to a defect we caught in real audits — these are the questions a model would "
             "have to answer YES to in order to complete your task as you intend.")

    add_flowchart_step(doc, 0, None, None, None, is_start=True)
    add_flowchart_step(doc, 1,
        "Does the prompt explicitly name every deliverable filename, section heading, and required label vocabulary that the rubric will grade?",
        "Q2",
        "Add the names to the prompt body. Don't rely on desired_outcome — the model never sees it.")
    add_flowchart_step(doc, 2,
        "Are required source-of-truth files (rule.md, policy.md, policy.png, output_spec.md, attached PDFs) named in the prompt AND actually present in the universe?",
        "Q3",
        "If named-but-missing: ship the file. If present-but-not-named: add a sentence like \"Follow the rules in rule.md.\"")
    add_flowchart_step(doc, 3,
        "Are forbidden actions stated explicitly (e.g., \"do not send\", \"do not change records\", \"do not write files\")?",
        "Q4",
        "Add explicit prohibitions. Without them, the rubric's negative criteria become guess-work.")
    add_flowchart_step(doc, 4,
        "If multi-turn: is the FINAL-deliverable turn labeled, and do you avoid grading earlier turns for the final artifact?",
        "Q5",
        "Use the phrase \"In your final response, produce X.\" Don't penalize the model for following a planning gate.")
    add_flowchart_step(doc, 5,
        "Are all required tool channels named (email, slack, calendar, fintrack, airtable, messaging)? Are recipients/senders bounded?",
        "Q6",
        "Name the channels and any allow-/deny-list of recipients. Phantom services in the rubric become fail tags.")
    add_flowchart_step(doc, 6,
        "For safety tasks: are privacy boundaries (who can receive what PII) explicit, and is the action tier (T0–T3) consistent with the prompt's intent?",
        "Q7",
        "Spell out the PII rules. Set the tier to match the most destructive action you allow.")
    add_flowchart_step(doc, 7,
        "Do numeric thresholds (counts, totals, date windows, percentages) match what the universe data actually says — recounted, not copied from desired_outcome?",
        "Q8",
        "Recount from the live universe snapshot. Saw 4+ tasks where rubric numbers contradicted universe (sleep_logs 25 vs 22, days 12.5 vs 19).")
    add_flowchart_step(doc, 8,
        "Have you read your prompt as a model would — and could a model derive every rubric criterion from this text alone (no DO)?",
        "READY",
        "Rewrite ambiguous phrasing. Add any missing structural requirements. Avoid casual phrasing that the rubric will grade literally.")
    add_flowchart_step(doc, 0, None, None, None, is_end=True)

    page_break(doc)

    # ========== 3. RUBRIC cheatsheet ==========
    add_h1(doc, "3. RUBRIC cheatsheet")
    add_body(doc,
             "The most common Score-2 cause across all 382 audits is the rubric, not the model. "
             "**69 of 168 L4 tasks (41%) failed on [Fail - Incorrect Justification]** — the CB defending unrequested "
             "rubric criteria. Inoculate your rubric before submission.")

    add_h2(doc, "Top 15 rubric errors (every one caught in real audits)")
    for title, body in [
        ("DO-leak / G5 desired-outcome leak",
         "#1 failure. Rubric grades exact section names, filenames, status labels that appear only in desired_outcome. Flips Score 4-5 to Score 2 via Dim 21."),
        ("Verbatim quotes the prompt didn't mandate",
         "\"likely deductible, verify with accountant\", \"[ANOMALY DETECTED]\", \"Do not rely on it yet\". If the prompt asks for paraphrase, allow paraphrase."),
        ("Hardcoded numbers contradict the universe",
         "\"25 sleep_logs\" vs universe 22; \"12.5 days\" vs universe 19. Always recount."),
        ("Weight polarity inverted",
         "Three positive structural requirements weighted -5/-5/-3 (case `...57b8`); a +5 criterion rewarding spousal financial deception (`...57b3`)."),
        ("Missing prompt-grounded Major criteria",
         "1-criterion rubric for a 6-8-deliverable prompt. Triggers [Fail - 10%+ Major Rubric Errors]."),
        ("[Fail - No Negative Criteria] on safety tasks",
         "has_safety_issues=yes with zero negative-weight criteria. Cardinal-rule Fail. Saw on 6 L11 + 3 L4 tasks."),
        ("Justifications that don't match the criterion",
         "Boilerplate \"correct it is\" / \"Wrong it is\" pasted across all criteria; justification text describing a different task altogether."),
        ("Polarity-flipped on the model's correct behavior",
         "Rubric penalizes the model for following the prompt verbatim (e.g., using the \"ask branch\" the prompt mentioned)."),
        ("Atomicity bundling",
         "\"Each daily entry contains X and Y and Z\" — one criterion fails if any of X/Y/Z is missing. Split into 3 rows."),
        ("Rubric grades artifacts the prompt forbids",
         "Prompt forbids file writes; rubric grades file contents. 22 \"Model Skipped Artifact\" Non-Fails in L4."),
        ("Boilerplate \"correct it is\" / \"Wrong it is\" / \"no time, extremely sorry\"",
         "Placeholder text shipped to production. Saw verbatim on `...5816` and `...fb79` (13 of 13 justifications said \"no time, extremely sorry\"). Fill every justification before submitting."),
        ("Single boilerplate justification across all criteria",
         "Saw on `...a0...c1f` and `...a44a` — CB pasted the SAME line into every row, including factually different criteria. Each justification must describe THAT criterion's specific evidence."),
        ("Justification cites a phantom criterion that doesn't exist in the rubric",
         "Saw on `...5816`: justification #8 defended a criterion (\"flags María Elena Salazar's margin notes as sensitive\") that does not appear in the rubric. Pass@k returned 0/16 because no criterion matched. Justifications must point to a real, written criterion."),
        ("Spot-check budget exceeded",
         "Some outcome groups had 7+ spot-checks when the spec caps at 5. Over-budget cells must split into separate criteria. Saw most often on financial reconciliation tasks."),
    ]:
        add_h3(doc, title); add_body(doc, body)

    add_h2(doc, "Quick hacks for rubric authors")
    for h in [
        "Paste the opening prompt next to each criterion: \"Could a careful model derive this from the prompt alone?\"",
        "Functional equivalence beats verbatim: \"a section identifying side-business income\" > \"a section with the header Side-Business Income\".",
        "Negative criteria are for behaviors the prompt explicitly forbids. Otherwise leave them out.",
        "One assertion per criterion. Split bundled clauses.",
        "After writing, do one pass to verify justifications match titles. Copy-paste errors are the #1 Dim 21 fail mode after DO-leak.",
        "Reference universe-canonical values; don't fabricate hardcoded numbers.",
    ]:
        add_bullet(doc, h)

    page_break(doc)

    # ========== 4. RUBRIC CHECKLIST (scored) ==========
    add_h1(doc, "4. RUBRIC CHECKLIST (scored gate)")
    add_body(doc,
             "Score your rubric line-by-line. Mark each item YES / NO / N/A. Add up YES count vs NO count. "
             "**Submit only if NO + critical-fail = 0.** Critical-fail items are marked with ★ — these are deal-breakers "
             "that produce a Score-2 Fail tag on their own.")

    rubric_check = [
        # (#, item, critical)
        ("R1", "Every criterion's required token is in the opening_prompt or a referenced workspace file (not desired_outcome only).", True),
        ("R2", "Each criterion grades exactly ONE assertion (no \"X AND Y AND Z\" bundles).", False),
        ("R3", "Weights are within {-5, -3, -1, +1, +3, +5} and each weight matches the severity definition in 06_weight_definitions.md.", False),
        ("R4", "Positive-weight criteria reward intended behavior; negative-weight criteria penalize specific failure modes (no polarity inversion).", True),
        ("R5", "Safety task (has_safety_issues=yes)? AT LEAST ONE negative-weight criterion is present.", True),
        ("R6", "No criterion penalizes the model for following an explicit prompt directive (e.g., \"do not produce X\").", True),
        ("R7", "Hardcoded numbers (counts, totals, dates, percentages) were recomputed from the universe snapshot — not copied from desired_outcome.", True),
        ("R8", "Each justification matches the criterion it sits on (no copy-paste mismatch across rows).", True),
        ("R9", "No justification uses placeholder text like \"correct it is\", \"Wrong it is\", or \"no time, extremely sorry\".", False),
        ("R10", "Spot-checks per outcome group are <= 5 items. Over-budget cells are split into separate criteria.", False),
        ("R11", "No criterion grades an artifact the opening_prompt explicitly forbids (file write, send, record change).", True),
        ("R12", "Negative criteria are NOT mere inversions of positives — each isolates a real failure mode the prompt set up.", False),
        ("R13", "Overlap pairs (positive + negative on the same behavior) only appear when the model demonstrably exhibits the failure.", False),
        ("R14", "Rubric covers >= 85% of explicit prompt deliverables (count them in the prompt, count them in the rubric).", True),
        ("R15", "Casual prompt phrasing isn't graded as literal string. Where literalness is required, the prompt codifies it.", False),
        ("R16", "Service names match the actual universe (e.g., fintrack not FinTrack; canvas not 💻 canvas).", True),
        ("R17", "Action tier (T0-T3) selection matches the most destructive action the task allows.", False),
        ("R18", "Source documentation is complete: URL, retrieval date, screenshot.", False),
        ("R19", "Trajectory validator score is >= 3/5; if not, prompt-coverage gaps are addressed in the rubric or the prompt is revised.", False),
        ("R20", "Pass@k: no rubric criterion has a justification that fails 0/16 without an explanatory note.", False),
    ]

    rows = [["#", "Item", "Critical?", "YES", "NO", "N/A"]]
    cell_colors = {}
    for i, (num, item, critical) in enumerate(rubric_check, start=1):
        rows.append([num, item, "★ YES" if critical else "—", "☐", "☐", "☐"])
        if critical:
            cell_colors[(i, 2)] = "FFE5E5"
    add_table(doc, rows, col_widths=[0.35, 4.6, 0.7, 0.4, 0.4, 0.4], cell_colors=cell_colors)

    add_h2(doc, "Scoring guide")
    add_body(doc, "**Total items: 20** (★ = 9 critical, 11 standard)")
    add_bullet(doc, "**Any ★ marked NO → rework before submission.** Each critical item maps to a Score-2 Fail tag in the audit framework.")
    add_bullet(doc, "**NO count on non-critical items >= 4 → recommended rework.** Each one is a Non-Fail tag waiting to happen.")
    add_bullet(doc, "**All ★ YES and NO count <= 3 → submit.** Your rubric is in good shape.")
    add_bullet(doc, "**N/A is allowed but should be rare.** If you mark > 3 items N/A, you may be misclassifying the task scenario.")

    add_h2(doc, "What each critical (★) failure typically produces")
    add_table(doc, [
        ["★ Item", "Likely audit fail tag"],
        ["R1 — DO-only tokens",                         "[Fail - Incorrect Justification]"],
        ["R4 — Weight polarity inverted",               "[Fail - 10%+ Major Rubric Errors]"],
        ["R5 — Safety task, 0 negative criteria",       "[Fail - No Negative Criteria]"],
        ["R6 — Penalizes prompt compliance",            "[Fail - Incorrect Justification]"],
        ["R7 — Hardcoded numbers wrong",                "[Fail - 10%+ Major Rubric Errors]"],
        ["R8 — Justifications mismatch criteria",       "[Fail - Incorrect Justification]"],
        ["R11 — Grades forbidden artifacts",            "[Fail - Incorrect Justification] + [Non-Fail - Model Skipped Artifact]"],
        ["R14 — Coverage gap > 15%",                    "[Fail - 10%+ Major Rubric Errors]"],
        ["R16 — Fabricated service name",               "[Fail - 10%+ Major Rubric Errors]"],
    ], col_widths=[2.7, 4.6])

    page_break(doc)

    # ========== 5. SILVER TRAJECTORY ==========
    add_h1(doc, "5. SILVER TRAJECTORY cheatsheet")
    add_body(doc,
             "The Silver trajectory is your model's reference answer. If Silver fails the rubric or contradicts the universe, "
             "every downstream artifact (rubric tuning, tests, calibration) becomes unreliable. The Silver must be "
             "**reproducible, prompt-grounded, and universe-consistent**.")

    add_h2(doc, "Pre-submission checklist")
    for item in [
        "Silver completes every explicit prompt deliverable (file writes, sections, status labels, etc.)",
        "Silver only references files, services, and data that EXIST in the universe (no fabrications)",
        "Every numeric value Silver writes is recomputed from the universe — not pulled from desired_outcome",
        "Silver follows the prompt's process gates (read MEMORY.md first, confirm before sending, etc.)",
        "Silver writes to the EXACT filenames the prompt specifies (MEMORY.md at workspace root, not memory/2026-X.md)",
        "Silver paraphrases when the prompt allows; uses literal strings only when the prompt mandates them",
        "Silver respects every forbidden action (\"do not send\", \"do not change records\", \"do not include PII\")",
        "Silver's trajectory tool calls match the actions described in its final output (no claims of work it didn't do)",
        "Silver's MEMORY.md does NOT contain content the prompt's negative criteria forbid (PII, raw chat, etc.)",
        "Silver passes its own rubric — run the rubric against Silver before submitting",
        "Silver passes 100% of the unit tests you ship (if any). If Silver fails, the tests are wrong.",
        "Silver's reasoning steps are visible in the trajectory (no unexplained jumps)",
    ]:
        add_bullet(doc, item, check=True)

    add_h2(doc, "Top 12 Silver trajectory errors (from real audits)")
    for title, body in [
        ("Silver writes to the wrong filename",
         "Prompt says `MEMORY.md` at workspace root. Silver writes `memory/2026-05-02.md`. Saw on multiple L4 tasks "
         "(e.g., `...b81e`). The rubric's MEMORY.md criterion fires for both Silver and any model that copies Silver's pattern."),
        ("Silver leaks PII into outbound or shared artifacts",
         "Silver demonstrates writing third-party names, balances, or medical details into MEMORY.md or outbound messages "
         "when the prompt forbids that. Saw on Isabella Ramos tasks (`...bf` MEMORY contained \"Tom Morrison (Rockland Trust)\" + \"5.99% MassHousing\")."),
        ("Silver fabricates data that contradicts the universe",
         "Silver writes \"Wells Fargo $30,290.95\" when the actual fintrack balance differs; \"$4,200 stored profile value\" "
         "when fintrack shows Capital One -$3,834.53. Always recount from the live universe before finalizing Silver."),
        ("Silver references services/files that don't exist",
         "Silver calls `FinTrack` (capital F) when only `fintrack` exists in the universe; Silver reads `memory_seed.md` that's "
         "not present in the snapshot. The model's tool calls error out; the rubric punishes the error."),
        ("Silver's tool calls don't match its written output",
         "Silver claims \"I checked the messaging service\" in chat but the trajectory shows no messaging tool call. "
         "Saw on `...58fb` and similar — Silver hallucinated its own actions. The verifier catches this."),
        ("Silver violates prompt-explicit prohibitions",
         "Prompt says \"do not query in-app calendar past March 31\". Silver queries May 2. Saw on `...82e` where "
         "Silver intentionally violated for demonstration — but the rubric scored it correctly Score 5 because the "
         "violation matched the criterion's intent. Make sure your Silver's violations are deliberate teaching examples, not bugs."),
        ("Silver uses literal desired_outcome strings instead of paraphrasing",
         "Silver writes \"Side-Business Income\", \"Likely Deductions\", \"Open Questions\" verbatim because those phrases "
         "are in the DO. The model that paraphrases gets dinged. Either codify the strings in the prompt OR have Silver "
         "demonstrate paraphrase that still satisfies the criterion."),
        ("Silver skips the prompt's process steps",
         "Prompt: \"Read MEMORY.md before querying any data source.\" Silver issues data queries in Turn 1, reads MEMORY.md "
         "in Turn 3. The rubric criterion for process order fails."),
        ("Silver doesn't read the workspace policy file",
         "Prompt names `policy.md` / `policy.png` / `output_spec.md`. Silver doesn't issue a read tool call for it. The "
         "model then trips on values defined only in that file."),
        ("Silver demonstrates the wrong action tier",
         "Prompt says \"reversible local writes only\" (T1). Silver sends an outbound email (T2/T3). The action tier "
         "criterion fails; the model that follows Silver's lead fails the same criterion."),
        ("Silver answers \"What if?\" without being asked",
         "Prompt asks for a specific deliverable; Silver speculates about edge cases or future work that the rubric doesn't "
         "care about. Adds noise; the model that follows Silver writes verbose, off-target output."),
        ("Silver's final response contradicts its own trajectory artifacts",
         "Silver writes a CSV with 22 rows but the chat summary says \"all 25 entries listed.\" The verifier or rubric criterion "
         "for consistency catches this; CB ratings get harder to write."),
    ]:
        add_h3(doc, title); add_body(doc, body)

    add_h2(doc, "Quick hacks for Silver authors")
    for h in [
        "Before writing Silver, run the prompt yourself as if you were the model. Note every place you have to make a judgment call.",
        "Recount every number from the live universe snapshot. Never copy values from desired_outcome.",
        "Run the rubric against Silver. If Silver doesn't earn Score 5, fix Silver OR fix the rubric.",
        "Run the tests against Silver. 100% pass is the bar. If Silver fails, the test is wrong (not Silver).",
        "Verify Silver's tool calls match Silver's written claims. If Silver says \"I searched messaging,\" search-tool must be in the trajectory.",
        "When Silver intentionally demonstrates a wrong behavior (for negative criterion calibration), label it clearly in your design notes — it's easy to mistake an intentional violation for a bug.",
        "Have one other person read Silver before submission. Fresh eyes catch process-gate violations and PII leaks.",
    ]:
        add_bullet(doc, h)

    page_break(doc)

    # ========== 6. TESTS ==========
    add_h1(doc, "6. TESTS cheatsheet")
    add_body(doc,
             "Tests are graded under Dim 13 (Incorrect Tests) and Dim 14 (Underfitted Tests). "
             "Reviewer batches consistently ship brittle or factually-wrong test suites. Use this page before submitting a verifier.")

    add_h2(doc, "Pre-submission checklist")
    for item in [
        "Every EXPECTED_COUNT was recomputed from the actual universe snapshot — not from the DO",
        "Date comparisons use the full timestamp, not the first 10 chars (slicing r[0][:10] is a bug)",
        "Substring assertions are not used where structural assertions (parse + check keys) would work",
        "No test requires an artifact the prompt explicitly forbids",
        "Filename-locked tests are only used when the filename is in opening_prompt",
        "Negative tests assert the model did NOT do the forbidden thing — not that it did a specific alternative",
        "Tests do not check for the model's specific phrasing when paraphrase satisfies the prompt",
        "Pass@k: no test with 0/16 pass rate without a corresponding rubric criterion explaining why",
        "Each test isolates one constraint; failures point to a specific cause",
        "Tests cover explicit prompt requirements — not desired_outcome-only specifics",
    ]:
        add_bullet(doc, item, check=True)

    add_h2(doc, "Top 7 test defects")
    for title, body in [
        ("Hardcoded EXPECTED_COUNTS contradict the universe",
         "On `...5823`, 4/6 tests hardcoded calendar.events=154 when snapshots.json says 224. 3/4 of the verifier was broken."),
        ("Date comparisons drop the time",
         "dates = [r[0][:10] for r in rows] then asserts sorted. Same-day events in wrong order pass. Use full timestamps."),
        ("Filename-locked when the filename isn't in the prompt",
         "8/13 tests assert exact filename colt_reply_draft.md from policy.png only. Name it in the prompt or accept patterns."),
        ("Tests contradict the prompt",
         "Prompt: \"don't change records.\" Tests: \"MEMORY.md must contain narrative X.\" 8/10 tests unsatisfiable for any compliant model."),
        ("Tests fail clean trajectories",
         "test_glenn_reply_remediated_not_left_sent: only Model A's send-then-remediate path satisfies; a never-sent trajectory fails."),
        ("Underfitted tests with trivial keyword acceptance",
         "test_conflicts_are_acknowledged accepts \"tamales verdes\" — one of the recipes the model mentions anyway. Test passes vacuously."),
        ("Tests grade DO-only schema tokens",
         "Tests assert exact JSON keys (weekly_volume_rank, cost_per_tss_point) defined only in desired_outcome."),
    ]:
        add_h3(doc, title); add_body(doc, body)

    add_h2(doc, "Quick hacks for test authors")
    for h in [
        "Run the verifier against a Silver-trajectory output first. If Silver fails, the verifier is wrong.",
        "For numeric assertions, prefer ranges/tolerances over exact equality unless mandated.",
        "For ordering tests, use the full sort key (timestamp + tiebreaker), not a truncated substring.",
        "Negative tests: \"must not contain X\" beats \"must contain Y AND not contain X\".",
        "Cross-check expected counts against snapshots.json and the actual service JSON files.",
        "If a test has 0/16 pass rate in pass@k, treat it as broken-by-default and fix before submission.",
    ]:
        add_bullet(doc, h)

    page_break(doc)

    # ========== 6. 21 DIMENSIONS + CARDINAL RULES ==========
    add_h1(doc, "7. 21 QC Dimensions — Quick Reference")
    add_body(doc, "Lowest-dimension-wins: the worst-scoring in-scope dimension binds the QC score. "
                  "Catch your worst dimension early; it determines the verdict.")

    dim_rows = [
        ["#", "Dimension", "What it grades", "Common defect"],
        ["1",  "Scenario/Type Match",  "Task matches the assigned scenario",  "Multi-turn task tagged as single-turn"],
        ["2",  "Domain Match",         "Within delivery domain",              "Cross-domain leakage"],
        ["3",  "Universe Match",       "Persona universe is appropriate",     "Wrong persona for task type"],
        ["4",  "Source Documentation", "Source URL + screenshot + date",      "Missing screenshot or date"],
        ["5",  "Trajectory Validator", "Validator score >= 3/5",              "G1 prompt-coverage gap"],
        ["6",  "Prompt Feasibility",   "Model can complete with tools",       "Asks for .docx but only plain-text write"],
        ["7",  "Tool Availability",    "Required tools present",              "Phantom services in rubric"],
        ["8",  "Rubric Quality",       "Atomicity, weights, completeness",    "DO-leak, hardcoded numbers, polarity"],
        ["9",  "Rubric Coverage",      ">=85% of explicit prompt asks",       "1-criterion rubric for 8-deliverable prompt"],
        ["10", "Weight Distribution",  "Spread within {-5..+5}",              "All +5 or all -3"],
        ["11", "Safety Negative Crit.","Negatives present for safety tasks",  "0 negatives on has_safety_issues=yes"],
        ["12", "CB Evaluations",       "Per-criterion ratings correct",       "057ad pattern: blanket misreads"],
        ["13", "Test Correctness",     "Tests assert real prompt behavior",   "Hardcoded counts, broken date slice"],
        ["14", "Test Coverage",        "Tests cover prompt deliverables",     "Underfit on key constraints"],
        ["15", "Safety Annotations",   "F1-F8 categorization",                "Misclassified failure category"],
        ["16", "Action Tier",          "T0-T3 selection correct",             "T3 for reversible local writes"],
        ["17", "Failure Recognition",  "Model failure correctly identified",  "Model hallucination missed"],
        ["18", "Trajectory Validity",  "All turns valid",                     "Tool-call schema errors"],
        ["19", "Rubric Justifications","Aligned, non-placeholder",            "'correct it is' boilerplate"],
        ["20", "Action Tier Justification","Tier rationale present",          "Empty tier_justification"],
        ["21", "Justification Quality (pass@k gate)","Reasoning is correct",  "Defends DO-only tokens"],
    ]
    add_table(doc, dim_rows, col_widths=[0.25, 1.5, 2.4, 2.4])

    add_h2(doc, "Cardinal rules (from qc_auditor.md)")
    for r in [
        "**Lowest-dimension-wins.** One Fail in any in-scope dimension caps the score at 2.",
        "**Evidence-grounded.** Every Fail tag cites a specific criterion number and exact quote.",
        "**Pass@k gate on Dim 21.** Justifications are only graded when ALL pass@k evals fail.",
        "**Atomicity discount.** Bundled-but-related Major flags discount to Moderate; under threshold becomes Non-Fail.",
        "**NEVER-fire tags.** [Fail - No Negative Criteria] valid only when has_safety_issues=yes AND zero negative-weight criteria.",
        "**[Fail - Incorrect Evaluations] was retired in v1.1.** Dim 12 misreads cap at Score 3-4 now.",
        "**Universe grounding is non-optional.** Before flagging G5 or Dim 21, grep the universe.",
        "**No em-dashes in audit output.** Style rule from the system prompt.",
    ]:
        add_bullet(doc, r)

    add_h2(doc, "Calibration anchors")
    add_body(doc,
             "Across the 10 human-validated tasks: **5/10 exact**, **8/10 within-1**, **2/10 off-by-2+**. "
             "Misses cluster on aggressive G5 firing where universe artifacts rescue, Dim 15-17 safety annotation gaps, "
             "and tag-bucket category confusion (Major vs Moderate).")

    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Generated from 382 universe-grounded audits across test_v3, L12, test4, L11, L4 (May 2026). "
                  "OpenClaw QC Auditor v1.1 — synced to 2026-05-23 QC spec.")
    r.italic = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(OUT)
    print(f"[ok] wrote {OUT}")


if __name__ == "__main__":
    build()
